import logging
import os
from pathlib import Path
import pandas as pd
import pdfplumber
from docx import Document
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Border, Side
from config import Config
'''
这段代码实现了将 PDF 和 Word 文档中的表格及文本提取并转换为 Excel 文件。
对于 PDF 文件，使用 pdfplumber 提取表格数据，自动合并跨页表格并格式化数值列，处理合并单元格；
对于 Word 文件，使用 python-docx 提取表格和段落内容，按格式填充 Excel，并对表格添加样式。
pdf_trans_excel 函数批量处理指定文件夹中的 PDF 和 Word 文件，将符合条件的文件转换为 Excel 文件，跳过 .doc 文件。
#注：效果一般，暂时弃用
'''
#todo：后续可以把不规则表格的处理加进来，比如不是没有单元格都有框线的表格
#todo：pdf 出错: single positional indexer is out-of-bounds需要解决

# 设置日志
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(message)s')
process_word = Config.process_word

def create_excel_writer(excel_path):
    """创建并返回一个 ExcelWriter 对象"""
    return pd.ExcelWriter(excel_path, engine='openpyxl')


def add_table_to_excel(df, writer, sheet_name, previous_table_header=None):
    """将 DataFrame 添加到 Excel 中，并处理标题行"""
    if df.columns[0] == 0 or all(isinstance(val, (int, float)) for val in df.iloc[0].values):
        # 如果当前表格没有标题行，将 previous_table_header 添加为标题
        df.columns = previous_table_header
    # 格式化数字列
    df = format_numbers(df)
    df.to_excel(writer, sheet_name=sheet_name, index=False)
    workbook = writer.book
    worksheet = workbook[sheet_name]
    split_merged_cells(worksheet)


def format_numbers(df):
    """将 DataFrame 中的数字列转换为数值类型"""
    for col in df.columns:
        if pd.api.types.is_numeric_dtype(df[col]):
            # 判断是否为整数或浮动数值
            df[col] = df[col].apply(lambda x: int(x) if x.is_integer() else round(x, 2))
    return df


def split_merged_cells(worksheet):
    """将合并的单元格拆分并填充数据"""
    # 获取所有的合并单元格
    merged_cells = worksheet.merged_cells.ranges
    for merged_cell in merged_cells:
        min_row, min_col, max_row, max_col = merged_cell.bounds
        value = worksheet.cell(min_row, min_col).value
        # 填充所有拆分后的单元格
        for row in range(min_row, max_row + 1):
            for col in range(min_col, max_col + 1):
                worksheet.cell(row, col, value=value)


def pdf_to_excel(pdf_path, excel_path):
    """使用 pdfplumber 将 PDF 表格提取到 Excel 文件中，并自动合并跨页表格"""
    all_tables = []
    previous_table_header = None
    combined_table = None
    with pdfplumber.open(pdf_path) as pdf:
        for page_num, page in enumerate(pdf.pages):
            try:
                tables = page.extract_tables()
                if tables:
                    for table in tables:
                        # 判断表格是否有效，避免空行或格式不一致
                        if not table or len(table) < 2:  # 至少有表头和一行数据
                            continue
                        has_header = bool(table[0])
                        current_df = pd.DataFrame(table[1:], columns=table[0]) if has_header else pd.DataFrame(table)
                        current_header = tuple(current_df.columns) if has_header else None

                        if previous_table_header == current_header or (not has_header and previous_table_header):
                            # 合并当前表格
                            combined_table = pd.concat([combined_table, current_df], ignore_index=True)
                        else:
                            # 保存上一张完整表格，并开始新表格
                            if combined_table is not None:
                                all_tables.append(combined_table)
                            combined_table = current_df
                            previous_table_header = current_header if has_header else previous_table_header
                    if page_num == len(pdf.pages) - 1 and combined_table is not None:
                        all_tables.append(combined_table)
            except Exception as e:
                logging.error(f"处理第 {page_num + 1} 页时出错: {e}")

    if not all_tables:
        logging.error(f"没有从 PDF 中提取到表格：{pdf_path}")
        return

    # 写入 Excel
    with create_excel_writer(excel_path) as writer:
        for i, df in enumerate(all_tables):
            sheet_name = f"Table_{i + 1}"
            add_table_to_excel(df, writer, sheet_name, previous_table_header)


def word_to_excel(doc_path, excel_path):
    """将 Word 文档中的表格和段落内容提取到 Excel 文件中"""
    doc = Document(doc_path)
    wb = Workbook()
    ws = wb.active
    ws.title = "WordContent"
    bold_font = Font(bold=True)
    fill_color = PatternFill(start_color="FFFF99", end_color="FFFF99", fill_type="solid")
    border_style = Border(left=Side(style="thin"), right=Side(style="thin"),
                          top=Side(style="thin"), bottom=Side(style="thin"))
    row_index = 1
    for table_num, table in enumerate(doc.tables, start=1):
        for row in table.rows:
            col_index = 1
            for cell in row.cells:
                cell_value = cell.text.strip()
                ws.cell(row=row_index, column=col_index, value=cell_value)
                ws.cell(row=row_index, column=col_index).border = border_style
                if table_num == 1:
                    ws.cell(row=row_index, column=col_index).fill = fill_color
                col_index += 1
            row_index += 1
        row_index += 1
    # 添加段落内容到新工作表
    ws_text = wb.create_sheet("TextContent")
    row_index = 1
    for para in doc.paragraphs:
        if para.text.strip():
            ws_text.cell(row=row_index, column=1, value=para.text.strip())
            row_index += 1
    wb.save(excel_path)


def report_transform_table(src_folder, process_word = process_word):
    """处理指定文件夹中的文件，将符合条件的 PDF 和（可选的）Word 文件转换为 Excel"""
    for root, _, files in os.walk(src_folder):
        for file_name in files:
            file_path = Path(root) / file_name
            file_stem = file_path.stem
            file_extension = file_path.suffix.lower()
            if file_extension == '.doc' and not process_word:
                logging.info(f"跳过 .doc 文件: {file_path}")
                continue
            if file_extension not in ['.docx', '.pdf']:
                continue
            if ((file_stem.endswith("表") and "报告" not in file_stem) or
                    any(keyword in file_stem for keyword in ["决算表", "预算表", "执行报表"])):
                excel_path = file_path.with_suffix('.xlsx')
                try:
                    # 根据文件类型调用相应的转换函数
                    if file_extension == '.pdf':
                        pdf_to_excel(file_path, excel_path)
                    elif file_extension == '.docx' and process_word:
                        word_to_excel(file_path, excel_path)
                    # file_path.unlink()  # 删除源文件
                    # logging.info(f"已处理并删除: {file_path}")
                except Exception as e:
                    logging.error(f"处理 {file_path} 出错: {e}")

# # 使用示例
# src_folder = r"H:\桌面\工作\已入职\项目内容\星火燎原\基于深度学习的自然场景文字信息提取模型\数据\训练集\安徽\安庆"
# report_trans_table(src_folder, process_word=True)
